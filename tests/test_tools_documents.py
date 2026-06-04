# tests/test_tools_documents.py
import pytest

from loom.tools.base import ToolError
from loom.tools.read import make_read_document
from loom.tools.web import WebSearchConfig, make_fetch_url


def test_read_document_xlsx(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Description", "Montant"])
    ws.append(["Prestation", 3600])
    wb.save(tmp_path / "facture.xlsx")
    out = make_read_document(str(tmp_path)).run({"path": "facture.xlsx"})
    assert "Description" in out and "3600" in out


def test_read_document_docx(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("Rapport trimestriel")
    d.add_paragraph("Conclusion : tout va bien.")
    d.save(tmp_path / "rapport.docx")
    out = make_read_document(str(tmp_path)).run({"path": "rapport.docx"})
    assert "Rapport trimestriel" in out and "Conclusion" in out


def test_read_document_rejects_text_format(tmp_path):
    (tmp_path / "note.txt").write_text("juste du texte", encoding="utf-8")
    with pytest.raises(ToolError, match="read_file"):
        make_read_document(str(tmp_path)).run({"path": "note.txt"})


def test_read_document_missing_file(tmp_path):
    with pytest.raises(ToolError, match="introuvable"):
        make_read_document(str(tmp_path)).run({"path": "absent.pdf"})


def test_fetch_url_returns_extracted_text(monkeypatch):
    import loom.tools.web as web

    # pas de vrai DNS en test : on neutralise le garde anti-SSRF et le fetch.
    monkeypatch.setattr(web, "_blocked_host_reason", lambda url: None)
    monkeypatch.setattr(
        web, "fetch_page", lambda url, cfg, snippet="": "Contenu de la page"
    )
    out = make_fetch_url(WebSearchConfig()).run({"url": "https://exemple.fr/a"})
    assert out == "Contenu de la page"


def test_fetch_url_rejects_non_http():
    with pytest.raises(ToolError):
        make_fetch_url(WebSearchConfig()).run({"url": "file:///etc/passwd"})


def test_fetch_url_blocks_internal_host_ssrf():
    # SSRF : une IP loopback/interne doit être refusée AVANT toute requête.
    with pytest.raises(ToolError, match="interne|interdit"):
        make_fetch_url(WebSearchConfig()).run({"url": "http://127.0.0.1:8080/admin"})


def test_fetch_url_blocks_cloud_metadata_ssrf():
    # 169.254.169.254 = métadonnées cloud (link-local) : interdit.
    with pytest.raises(ToolError, match="interne|interdit"):
        make_fetch_url(WebSearchConfig()).run(
            {"url": "http://169.254.169.254/latest/meta-data/"}
        )
